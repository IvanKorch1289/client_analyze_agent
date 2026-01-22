"""
Export routes for report downloads.

Provides endpoints for:
- Excel export (.xlsx)
- Word export (.docx)
- History timeline
"""

from datetime import datetime
from io import BytesIO
from typing import Any, Dict, List, Optional

from fastapi import APIRouter, HTTPException, Query
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from app.utility.logging_client import logger

export_router = APIRouter(prefix="/export", tags=["export"])


# ============================================================================
# SCHEMAS
# ============================================================================


class AnalysisSnapshot(BaseModel):
    """Snapshot of analysis for timeline."""

    report_id: str
    date: str
    risk_score: int
    risk_level: str
    factors_count: int
    score_change: Optional[int] = None
    new_factors: List[str] = []
    removed_factors: List[str] = []


# ============================================================================
# EXPORT ENDPOINTS
# ============================================================================


@export_router.get("/{report_id}/excel")
async def export_to_excel(report_id: str) -> StreamingResponse:
    """
    Export report to Excel format (.xlsx).

    Creates a multi-sheet workbook with:
    - Main info (company, risk score, date)
    - Risk factors
    - Recommendations
    - Findings
    """
    try:
        import pandas as pd
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="pandas and openpyxl are required for Excel export. Install with: pip install pandas openpyxl",
        )

    # Get report from storage
    report = await _get_report(report_id)
    if not report:
        raise HTTPException(status_code=404, detail=f"Report {report_id} not found")

    try:
        output = BytesIO()
        with pd.ExcelWriter(output, engine="openpyxl") as writer:
            # Sheet 1: Main information
            metadata = report.get("metadata", {})
            risk = report.get("risk_assessment", {})

            df_main = pd.DataFrame(
                [
                    {
                        "Компания": metadata.get("client_name", "N/A"),
                        "ИНН": metadata.get("inn", "N/A"),
                        "Риск-скор": risk.get("score", 0),
                        "Уровень риска": risk.get("level", "unknown"),
                        "Дата анализа": metadata.get("analysis_date", ""),
                        "Источников данных": metadata.get("data_sources_count", 0),
                        "Успешных источников": metadata.get("successful_sources", 0),
                    }
                ]
            )
            df_main.to_excel(writer, sheet_name="Основное", index=False)

            # Sheet 2: Risk factors
            factors = risk.get("factors", [])
            if factors:
                df_factors = pd.DataFrame({"Фактор риска": factors})
                df_factors.to_excel(writer, sheet_name="Факторы риска", index=False)

            # Sheet 3: Recommendations
            recommendations = report.get("recommendations", [])
            if recommendations:
                df_rec = pd.DataFrame({"Рекомендация": recommendations})
                df_rec.to_excel(writer, sheet_name="Рекомендации", index=False)

            # Sheet 4: Findings
            findings = report.get("findings", [])
            if findings:
                df_findings = pd.DataFrame(findings)
                df_findings.to_excel(writer, sheet_name="Находки", index=False)

            # Sheet 5: Risk categories (if available)
            categories = risk.get("categories", {})
            if categories:
                df_categories = pd.DataFrame([{"Категория": k, "Балл": v} for k, v in categories.items()])
                df_categories.to_excel(writer, sheet_name="Категории рисков", index=False)

        output.seek(0)

        filename = f"report_{report_id}_{datetime.now().strftime('%Y%m%d')}.xlsx"
        logger.info(f"Excel export generated: {filename}", component="export")

        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    except Exception as e:
        logger.error(f"Excel export failed: {e}", component="export")
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")


@export_router.get("/{report_id}/word")
async def export_to_word(report_id: str) -> StreamingResponse:
    """
    Export report to Word format (.docx).

    Creates a formatted document with:
    - Company header
    - Risk assessment summary
    - Detailed findings
    - Recommendations
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="python-docx is required for Word export. Install with: pip install python-docx",
        )

    # Get report from storage
    report = await _get_report(report_id)
    if not report:
        raise HTTPException(status_code=404, detail=f"Report {report_id} not found")

    try:
        doc = Document()

        # Title
        metadata = report.get("metadata", {})
        title = doc.add_heading(f"Отчёт по анализу: {metadata.get('client_name', 'N/A')}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Basic info
        doc.add_paragraph(f"ИНН: {metadata.get('inn', 'Не указан')}")
        doc.add_paragraph(f"Дата анализа: {metadata.get('analysis_date', 'N/A')}")
        doc.add_paragraph()

        # Risk assessment
        risk = report.get("risk_assessment", {})
        doc.add_heading("Оценка рисков", level=1)

        risk_para = doc.add_paragraph()
        risk_para.add_run("Риск-скор: ").bold = True
        risk_para.add_run(f"{risk.get('score', 0)}/100")

        level_para = doc.add_paragraph()
        level_para.add_run("Уровень: ").bold = True
        level_para.add_run(risk.get("level", "unknown").upper())

        # Risk factors
        factors = risk.get("factors", [])
        if factors:
            doc.add_heading("Факторы риска", level=2)
            for factor in factors:
                doc.add_paragraph(factor, style="List Bullet")

        # Summary
        summary = report.get("summary", "")
        if summary:
            doc.add_heading("Резюме", level=1)
            doc.add_paragraph(summary)

        # Findings
        findings = report.get("findings", [])
        if findings:
            doc.add_heading("Детальные находки", level=1)
            for finding in findings:
                if isinstance(finding, dict):
                    p = doc.add_paragraph()
                    category = finding.get("category", "")
                    severity = finding.get("severity", "")
                    desc = finding.get("description", str(finding))
                    p.add_run(f"[{category}] ").bold = True
                    if severity:
                        p.add_run(f"({severity}) ")
                    p.add_run(desc)
                else:
                    doc.add_paragraph(str(finding), style="List Bullet")

        # Recommendations
        recommendations = report.get("recommendations", [])
        if recommendations:
            doc.add_heading("Рекомендации", level=1)
            for i, rec in enumerate(recommendations, 1):
                doc.add_paragraph(rec, style="List Number")

        # Reasoning trace (if available)
        reasoning = report.get("reasoning", {})
        if reasoning:
            doc.add_heading("Обоснование оценки (Chain-of-Thought)", level=1)

            calculation = reasoning.get("calculation", "")
            if calculation:
                doc.add_paragraph(f"Расчёт: {calculation}")

            confidence = reasoning.get("confidence", 0)
            if confidence:
                doc.add_paragraph(f"Уверенность: {confidence * 100:.0f}%")

        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.add_run(
            f"Отчёт сгенерирован системой Client Analysis Agent | {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        ).italic = True

        # Save to buffer
        output = BytesIO()
        doc.save(output)
        output.seek(0)

        filename = f"report_{report_id}_{datetime.now().strftime('%Y%m%d')}.docx"
        logger.info(f"Word export generated: {filename}", component="export")

        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    except Exception as e:
        logger.error(f"Word export failed: {e}", component="export")
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")


# ============================================================================
# HISTORY ENDPOINTS
# ============================================================================


@export_router.get("/history/{inn}")
async def get_analysis_timeline(
    inn: str,
    limit: int = Query(default=20, le=100, description="Maximum number of records"),
) -> List[AnalysisSnapshot]:
    """
    Get analysis history for a company by INN.

    Returns a timeline showing:
    - All previous analyses
    - Risk score changes over time
    - New/removed risk factors between analyses
    """
    try:
        from app.storage.tarantool import TarantoolClient

        client = await TarantoolClient.get_instance()
        repo = client.get_reports_repository()
        reports = await repo.get_reports_by_inn(inn, limit=limit)
    except Exception as e:
        logger.error(f"Failed to get history: {e}", component="export")
        raise HTTPException(status_code=500, detail=f"Failed to retrieve history: {str(e)}")

    if not reports:
        return []

    # Sort by date
    sorted_reports = sorted(reports, key=lambda r: r.get("created_at", ""))

    timeline = []
    prev_report = None

    for report in sorted_reports:
        risk = report.get("risk_assessment", {})
        factors = risk.get("factors", [])

        snapshot = AnalysisSnapshot(
            report_id=report.get("id", ""),
            date=report.get("created_at", ""),
            risk_score=risk.get("score", 0),
            risk_level=risk.get("level", "unknown"),
            factors_count=len(factors),
        )

        if prev_report:
            prev_risk = prev_report.get("risk_assessment", {})
            prev_factors = set(prev_risk.get("factors", []))
            curr_factors = set(factors)

            snapshot.score_change = risk.get("score", 0) - prev_risk.get("score", 0)
            snapshot.new_factors = list(curr_factors - prev_factors)
            snapshot.removed_factors = list(prev_factors - curr_factors)

        timeline.append(snapshot)
        prev_report = report

    logger.info(
        f"History timeline retrieved: {len(timeline)} records for INN {inn}",
        component="export",
    )

    return timeline


# ============================================================================
# HELPER FUNCTIONS
# ============================================================================


async def _get_report(report_id: str) -> Optional[Dict[str, Any]]:
    """Get report from storage by ID."""
    try:
        from app.storage.tarantool import TarantoolClient

        client = await TarantoolClient.get_instance()
        repo = client.get_reports_repository()
        return await repo.get(report_id)
    except Exception as e:
        logger.error(f"Failed to get report {report_id}: {e}", component="export")
        return None


__all__ = ["export_router"]
